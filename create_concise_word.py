"""
Create Word Document - BTP Report Concise (Line-by-line same as PDF)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_word_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # ============================================
    # TITLE PAGE (Matching uploaded picture format)
    # ============================================
    
    doc.add_paragraph()
    
    # Main Title - COMMODITY PRICE PREDICTION SYSTEM
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('COMMODITY PRICE PREDICTION SYSTEM')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Using Machine Learning and Deep Learning')
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Report type
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bachelor's Thesis Project Report (BTP1)")
    run.font.size = Pt(12)
    run.font.italic = True
    
    doc.add_paragraph()
    
    # Submitted by
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Submitted by:')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Gaurav Kumar(22AG36012)')
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Supervisor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Under the supervision of:')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Dr. Priyabrata Pradhan')
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Department
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Agriculture and Food Engineering')
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('IIT Kharagpur')
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Institution name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('INDIAN INSTITUTE OF TECHNOLOGY')
    run.font.size = Pt(16)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Kharagpur')
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[IIT Kharagpur Logo]')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Department footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Agricultural and Food Engineering Department')
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Indian Institute of Technology Kharagpur')
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Semester and Date (blue)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Autumn Semester, 2025-26')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('November 27, 2025')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    doc.add_page_break()
    
    # ============================================
    # CERTIFICATE
    # ============================================
    
    doc.add_heading('Certificate', level=1)
    
    doc.add_paragraph('This is to certify that the project titled "Commodity Price Prediction Using Machine Learning and Deep Learning" submitted by Gaurav Kumar (22AG36012) to IIT Kharagpur is a record of bonafide work carried out under my supervision.')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Dr. Priyabrata Pradhan\n').bold = True
    p.add_run('Department of Agriculture and Food Engineering, IIT Kharagpur\nDate: November 27, 2025')
    
    doc.add_page_break()
    
    # ============================================
    # ABSTRACT
    # ============================================
    
    doc.add_heading('Abstract', level=1)
    
    doc.add_paragraph('This project develops a machine learning system for predicting agricultural commodity prices in West Bengal, India. Using historical data from 2014-2025 comprising 177,320 records across 18 districts and 61 markets, we implement two models: XGBoost and Deep Neural Network.')
    
    p = doc.add_paragraph()
    p.add_run('The XGBoost model achieves ').font.name = 'Times New Roman'
    p.add_run('5.43% MAPE').bold = True
    p.add_run(' with ')
    p.add_run('R² = 0.9453').bold = True
    p.add_run(', while the Neural Network achieves ')
    p.add_run('4.64% MAPE').bold = True
    p.add_run(' with ')
    p.add_run('R² = 0.9601').bold = True
    p.add_run('. For 2025 predictions, the Neural Network achieves excellent ')
    p.add_run('4.44% MAPE').bold = True
    p.add_run(' with ')
    p.add_run('R² = 0.9724').bold = True
    p.add_run('. The system predicts prices for Rice, Jute, and Wheat with 7-day forecasting capability.')
    
    doc.add_paragraph('A web application built with Flask and React provides an accessible interface for stakeholders to obtain price predictions.')
    
    p = doc.add_paragraph()
    p.add_run('Keywords: ').bold = True
    p.add_run('Machine Learning, XGBoost, Neural Networks, Price Prediction, Agriculture')
    
    doc.add_page_break()
    
    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    
    doc.add_heading('Contents', level=1)
    
    toc = [
        ('Certificate', '1'),
        ('Abstract', '2'),
        ('1 Introduction', '4'),
        ('  1.1 Background', '4'),
        ('  1.2 Problem Statement', '4'),
        ('  1.3 Objectives', '4'),
        ('2 Methodology', '5'),
        ('  2.1 Dataset', '5'),
        ('  2.2 Feature Engineering', '6'),
        ('  2.3 Models', '6'),
        ('3 Results and Analysis', '8'),
        ('  3.1 Model Performance', '8'),
        ('  3.2 Commodity-wise Performance', '8'),
        ('  3.3 Discussion', '9'),
        ('4 System Implementation', '10'),
        ('5 Conclusion', '12'),
        ('References', '14'),
    ]
    
    for item, page in toc:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run('\t\t' + page)
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 1: INTRODUCTION
    # ============================================
    
    doc.add_heading('Chapter 1: Introduction', level=1)
    
    doc.add_heading('1.1 Background', level=2)
    doc.add_paragraph("Agriculture employs over 50% of India's workforce. Price volatility significantly impacts farmers' livelihoods, leading to distress selling and financial instability. West Bengal, a major agricultural state, produces significant quantities of rice, jute, and wheat with complex pricing patterns.")
    
    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph('Develop an accurate commodity price prediction system that:')
    doc.add_paragraph('1. Predicts prices for Rice, Jute, and Wheat in West Bengal', style='List Number')
    doc.add_paragraph('2. Provides 7-day ahead forecasts', style='List Number')
    doc.add_paragraph('3. Incorporates weather, economic indicators, and historical patterns', style='List Number')
    doc.add_paragraph('4. Offers an accessible web interface', style='List Number')
    
    doc.add_heading('1.3 Objectives', level=2)
    doc.add_paragraph('1. Compile comprehensive historical price data (2014-2025)', style='List Number')
    doc.add_paragraph('2. Design relevant features capturing temporal and economic patterns', style='List Number')
    doc.add_paragraph('3. Implement XGBoost and Neural Network models', style='List Number')
    doc.add_paragraph('4. Develop a production-ready web application', style='List Number')
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 2: METHODOLOGY
    # ============================================
    
    doc.add_heading('Chapter 2: Methodology', level=1)
    
    doc.add_heading('2.1 Dataset', level=2)
    
    # Dataset Overview Table
    doc.add_paragraph('Table 2.1: Dataset Overview')
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ('Parameter', 'Value'),
        ('Total Records', '177,320'),
        ('Time Period', '2014-2025 (11 years)'),
        ('Districts', '18'),
        ('Markets', '61'),
        ('Commodities', 'Rice, Jute, Wheat'),
        ('Database Size', '51.14 MB'),
    ]
    
    for i, (param, value) in enumerate(data):
        table.rows[i].cells[0].text = param
        table.rows[i].cells[1].text = value
        if i == 0:
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Commodity Distribution Table
    doc.add_paragraph('Table 2.2: Commodity Distribution')
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'
    
    comm_data = [
        ('Commodity', 'Records', '%'),
        ('Rice', '130,572', '75.4%'),
        ('Jute', '34,425', '19.9%'),
        ('Wheat', '8,097', '4.7%'),
    ]
    
    for i, row_data in enumerate(comm_data):
        for j, cell_data in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_data
            if i == 0:
                table2.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('Data sources include Agmarknet (prices), IMD (weather), RBI (economic indicators), and Ministry of Agriculture (MSP, production data).')
    
    doc.add_heading('2.2 Feature Engineering', level=2)
    doc.add_paragraph('We designed 36 features in the following categories:')
    
    p = doc.add_paragraph()
    p.add_run('Temporal Features: ').bold = True
    p.add_run('Year, month, day, quarter, day of week, weekend indicator, seasonal flags (monsoon, winter, summer).')
    
    p = doc.add_paragraph()
    p.add_run('Categorical Features: ').bold = True
    p.add_run('District, market, commodity, and variety (label encoded).')
    
    p = doc.add_paragraph()
    p.add_run('Economic Indicators: ').bold = True
    p.add_run('CPI, per capita income, food subsidy, MSP (Minimum Support Price).')
    
    p = doc.add_paragraph()
    p.add_run('Agricultural Parameters: ').bold = True
    p.add_run('Temperature, rainfall, area, production, yield, fertilizer consumption, export/import data.')
    
    p = doc.add_paragraph()
    p.add_run('Price Statistics: ').bold = True
    p.add_run('Commodity average, market average, district-commodity average, monthly averages.')
    
    doc.add_heading('2.3 Models', level=2)
    
    p = doc.add_paragraph()
    p.add_run('XGBoost (Extreme Gradient Boosting)').bold = True
    doc.add_paragraph('XGBoost builds an ensemble of decision trees sequentially, optimizing a regularized objective function.')
    doc.add_paragraph('Key Hyperparameters: 1000 estimators, max depth 8, learning rate 0.05, L1/L2 regularization, GPU acceleration.')
    
    p = doc.add_paragraph()
    p.add_run('Neural Network').bold = True
    doc.add_paragraph('A 5-layer deep neural network with architecture:')
    doc.add_paragraph('• Input: 36 features')
    doc.add_paragraph('• Hidden layers: 256 → 128 → 64 → 32 → 16 neurons')
    doc.add_paragraph('• Activation: ReLU with BatchNorm and Dropout')
    doc.add_paragraph('• Output: 1 neuron (price prediction)')
    doc.add_paragraph('• Optimizer: Adam (learning rate 0.001)')
    doc.add_paragraph('• Early stopping with patience 20')
    
    doc.add_heading('2.4 Evaluation Metrics', level=2)
    doc.add_paragraph('• MAE: Mean Absolute Error (Rs)')
    doc.add_paragraph('• RMSE: Root Mean Squared Error (Rs)')
    doc.add_paragraph('• MAPE: Mean Absolute Percentage Error (%)')
    doc.add_paragraph('• R²: Coefficient of Determination')
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 3: RESULTS
    # ============================================
    
    doc.add_heading('Chapter 3: Results and Analysis', level=1)
    
    doc.add_heading('3.1 Model Performance', level=2)
    
    # Performance Table
    doc.add_paragraph('Table 3.1: Model Performance Comparison')
    perf_table = doc.add_table(rows=5, cols=3)
    perf_table.style = 'Table Grid'
    
    perf_data = [
        ('Metric', 'XGBoost', 'Neural Network'),
        ('MAE (Rs)', '167.42', '143.73'),
        ('RMSE (Rs)', '285.36', '253.66'),
        ('MAPE (%)', '5.43', '4.64'),
        ('R² Score', '0.9453', '0.9601'),
    ]
    
    for i, row_data in enumerate(perf_data):
        for j, cell_data in enumerate(row_data):
            perf_table.rows[i].cells[j].text = cell_data
            if i == 0:
                perf_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
            if j == 2 and i > 0:  # Bold NN values as they are better
                perf_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Accuracy Distribution Table
    doc.add_paragraph('Table 3.2: Prediction Accuracy Distribution')
    acc_table = doc.add_table(rows=4, cols=3)
    acc_table.style = 'Table Grid'
    
    acc_data = [
        ('Error Threshold', 'XGBoost', 'Neural Network'),
        ('Within 5%', '89.2%', '92.4%'),
        ('Within 10%', '97.5%', '98.1%'),
        ('Within 15%', '99.2%', '99.5%'),
    ]
    
    for i, row_data in enumerate(acc_data):
        for j, cell_data in enumerate(row_data):
            acc_table.rows[i].cells[j].text = cell_data
            if i == 0:
                acc_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Commodity-wise Performance', level=2)
    
    doc.add_paragraph('Table 3.3: MAPE by Commodity (%)')
    comm_perf = doc.add_table(rows=4, cols=3)
    comm_perf.style = 'Table Grid'
    
    comm_perf_data = [
        ('Commodity', 'XGBoost', 'Neural Network'),
        ('Rice', '5.21', '4.42'),
        ('Jute', '5.89', '4.95'),
        ('Wheat', '5.45', '4.68'),
    ]
    
    for i, row_data in enumerate(comm_perf_data):
        for j, cell_data in enumerate(row_data):
            comm_perf.rows[i].cells[j].text = cell_data
            if i == 0:
                comm_perf.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Validation Results
    doc.add_paragraph('Table 3.4: Sample Predictions vs Actual Prices (2025 Data)')
    val_table = doc.add_table(rows=4, cols=5)
    val_table.style = 'Table Grid'
    
    val_data = [
        ('Commodity', 'District', 'Actual', 'XGBoost', 'NN'),
        ('Wheat', 'North 24 Parganas', 'Rs 2,060', 'Rs 2,180', 'Rs 2,105'),
        ('Jute', 'Murshidabad', 'Rs 5,880', 'Rs 6,120', 'Rs 5,894'),
        ('Rice', 'Medinipur(W)', 'Rs 3,800', 'Rs 3,950', 'Rs 3,811'),
    ]
    
    for i, row_data in enumerate(val_data):
        for j, cell_data in enumerate(row_data):
            val_table.rows[i].cells[j].text = cell_data
            if i == 0:
                val_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 Feature Importance', level=2)
    doc.add_paragraph('Top features influencing predictions (XGBoost):')
    doc.add_paragraph('1. Commodity average price (18.7%)')
    doc.add_paragraph('2. Market average price (15.6%)')
    doc.add_paragraph('3. MSP - Minimum Support Price (13.4%)')
    doc.add_paragraph('4. Variety average price (9.8%)')
    doc.add_paragraph('5. CPI (8.7%)')
    
    doc.add_heading('3.4 Discussion', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Neural Network Advantages:').bold = True
    doc.add_paragraph('• Superior accuracy (4.64% MAPE vs 5.43%)')
    doc.add_paragraph('• Excellent 2025 predictions (4.44% MAPE, R² = 0.9724)')
    doc.add_paragraph('• Effectively captures temporal and seasonal patterns')
    doc.add_paragraph('• Better generalization on unseen data')
    
    p = doc.add_paragraph()
    p.add_run('XGBoost Observations:').bold = True
    doc.add_paragraph('• Good performance (5.43% MAPE)')
    doc.add_paragraph('• Handles categorical features effectively')
    doc.add_paragraph('• Provides feature importance insights')
    doc.add_paragraph('• Useful as fallback model')
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 4: SYSTEM IMPLEMENTATION
    # ============================================
    
    doc.add_heading('Chapter 4: System Implementation', level=1)
    
    doc.add_heading('4.1 Architecture', level=2)
    doc.add_paragraph('The system follows a three-tier architecture:')
    doc.add_paragraph('• Presentation Layer: React.js frontend')
    doc.add_paragraph('• Application Layer: Flask REST API')
    doc.add_paragraph('• Data Layer: SQLite database + trained models')
    
    doc.add_heading('4.2 Technology Stack', level=2)
    
    doc.add_paragraph('Table 4.1: Technology Stack')
    tech_table = doc.add_table(rows=6, cols=2)
    tech_table.style = 'Table Grid'
    
    tech_data = [
        ('Component', 'Technology'),
        ('Frontend', 'React.js 18'),
        ('Backend', 'Flask 3.1, Waitress WSGI'),
        ('ML Framework', 'XGBoost 3.1.2, TensorFlow 2.20'),
        ('Database', 'SQLite'),
        ('Language', 'Python 3.10'),
    ]
    
    for i, row_data in enumerate(tech_data):
        for j, cell_data in enumerate(row_data):
            tech_table.rows[i].cells[j].text = cell_data
            if i == 0:
                tech_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('4.3 Web Application Features', level=2)
    doc.add_paragraph('1. Model Selection: Choose between XGBoost or Neural Network')
    doc.add_paragraph('2. Input Form: Select commodity, district, market, variety, date')
    doc.add_paragraph('3. 7-Day Forecast: Display predictions for current day + 6 days')
    doc.add_paragraph('4. Responsive Design: Works on desktop and mobile devices')
    
    doc.add_heading('4.4 API Endpoints', level=2)
    
    doc.add_paragraph('Table 4.2: REST API')
    api_table = doc.add_table(rows=5, cols=3)
    api_table.style = 'Table Grid'
    
    api_data = [
        ('Endpoint', 'Method', 'Description'),
        ('/', 'GET', 'Main page'),
        ('/predict', 'POST', 'Get price predictions'),
        ('/get_markets', 'POST', 'Markets for district'),
        ('/get_varieties', 'POST', 'Varieties for commodity'),
    ]
    
    for i, row_data in enumerate(api_data):
        for j, cell_data in enumerate(row_data):
            api_table.rows[i].cells[j].text = cell_data
            if i == 0:
                api_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 5: CONCLUSION
    # ============================================
    
    doc.add_heading('Chapter 5: Conclusion', level=1)
    
    doc.add_heading('5.1 Summary', level=2)
    doc.add_paragraph('This project successfully developed a commodity price prediction system achieving:')
    doc.add_paragraph('• Neural Network with 4.64% MAPE and R² = 0.9601')
    doc.add_paragraph('• XGBoost model with 5.43% MAPE and R² = 0.9453')
    doc.add_paragraph('• 2025 predictions: Neural Network achieves 4.44% MAPE, R² = 0.9724')
    doc.add_paragraph('• Coverage of 18 districts, 61 markets, 3 commodities')
    doc.add_paragraph('• Production-ready web application')
    
    doc.add_heading('5.2 Contributions', level=2)
    doc.add_paragraph('1. Novel feature set combining economic and agricultural indicators')
    doc.add_paragraph('2. Comparative analysis of gradient boosting vs deep learning')
    doc.add_paragraph('3. Deployed system accessible to farmers and policymakers')
    
    doc.add_heading('5.3 Limitations', level=2)
    doc.add_paragraph('• Limited to West Bengal region')
    doc.add_paragraph('• Does not account for sudden policy changes or disasters')
    doc.add_paragraph('• Requires continuous data updates')
    
    doc.add_heading('5.4 Future Work', level=2)
    doc.add_paragraph('1. Implement LSTM for better temporal pattern capture')
    doc.add_paragraph('2. Extend coverage to pan-India')
    doc.add_paragraph('3. Add more commodities (vegetables, pulses)')
    doc.add_paragraph('4. Develop mobile application')
    doc.add_paragraph('5. Integrate real-time weather data APIs')
    
    doc.add_page_break()
    
    # ============================================
    # REFERENCES
    # ============================================
    
    doc.add_heading('References', level=1)
    
    refs = [
        '[1] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. KDD.',
        '[2] Hochreiter, S., & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation.',
        '[3] Srivastava, N. et al. (2014). Dropout: Preventing Overfitting. JMLR.',
        '[4] Kingma, D. P., & Ba, J. (2014). Adam Optimizer. arXiv:1412.6980.',
        '[5] Breiman, L. (2001). Random Forests. Machine Learning.',
        '[6] Abadi, M. et al. (2016). TensorFlow. OSDI.',
        '[7] Pedregosa, F. et al. (2011). Scikit-learn. JMLR.',
    ]
    
    for ref in refs:
        doc.add_paragraph(ref)
    
    # Save
    doc.save('BTP_Report_Concise.docx')
    print('✅ Word document saved: BTP_Report_Concise.docx')

if __name__ == '__main__':
    create_word_report()
