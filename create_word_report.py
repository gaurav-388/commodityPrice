"""
Create Word Document from BTP Report
Uses python-docx to create a Word document version of the BTP Report
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

def create_word_report():
    """Create comprehensive BTP Report in Word format"""
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # ============================================
    # TITLE PAGE
    # ============================================
    
    # Add logo placeholder
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('[IIT Kharagpur Logo]')
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Institution
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inst.add_run('INDIAN INSTITUTE OF TECHNOLOGY KHARAGPUR')
    run.font.size = Pt(16)
    run.font.bold = True
    
    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dept.add_run('Department of Agriculture and Food Engineering')
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('COMMODITY PRICE PREDICTION SYSTEM')
    run.font.size = Pt(20)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Using Machine Learning and Deep Learning')
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    project_type = doc.add_paragraph()
    project_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project_type.add_run('Bachelor of Technology Project (BTP) Report')
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student info
    student = doc.add_paragraph()
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student.add_run('Submitted by:\nGaurav Kumar\nB.Tech, Agricultural and Food Engineering\nIIT Kharagpur')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Supervisor info
    supervisor = doc.add_paragraph()
    supervisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = supervisor.add_run('Under the supervision of:\nDr. Prasun Kumar Pany\nDepartment of Agriculture and Food Engineering\nIIT Kharagpur')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('Academic Year: 2024-2025\nNovember 2025')
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ============================================
    # ABSTRACT
    # ============================================
    
    doc.add_heading('Abstract', level=1)
    
    abstract_text = """Agricultural commodity price prediction is a critical challenge in developing economies like India, where price volatility significantly impacts farmers' livelihoods and food security. This project presents a comprehensive machine learning-based system for predicting commodity prices in West Bengal, India, utilizing historical price data spanning from 2014 to 2025.

We develop and compare two predictive models: XGBoost (Extreme Gradient Boosting) and a Deep Neural Network with Backpropagation. The system incorporates 30 carefully engineered features including temporal patterns, economic indicators (CPI, MSP, food subsidies), agricultural parameters (temperature, rainfall, fertilizer consumption), and market-specific statistics.

Our XGBoost model achieves a Mean Absolute Percentage Error (MAPE) of 5.43% with an R² score of 0.9453, while the Neural Network model achieves a MAPE of 4.64% with an R² score of 0.9601. For 2025 predictions specifically, the Neural Network achieves an excellent MAPE of 4.44% with R² = 0.9724. The system successfully predicts prices for three major commodities—Rice, Jute, and Wheat—across 18 districts and 61 markets in West Bengal.

A web-based application built using Flask and React provides an intuitive interface for stakeholders to access 7-day price forecasts. The system is deployed using a production-grade Waitress WSGI server ensuring reliability and scalability.

Keywords: Machine Learning, XGBoost, Neural Networks, Commodity Price Prediction, Agricultural Analytics, Time Series Forecasting, Deep Learning"""
    
    doc.add_paragraph(abstract_text)
    doc.add_page_break()
    
    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1. Introduction', 11),
        ('   1.1 Background and Motivation', 11),
        ('   1.2 Problem Statement', 11),
        ('   1.3 Objectives', 12),
        ('2. Literature Review', 14),
        ('   2.1 Traditional Approaches', 14),
        ('   2.2 Machine Learning Approaches', 15),
        ('   2.3 Deep Learning Approaches', 16),
        ('3. Methodology', 17),
        ('   3.1 Data Collection', 17),
        ('   3.2 Feature Engineering', 18),
        ('   3.3 Model Architecture', 20),
        ('4. System Design and Implementation', 26),
        ('   4.1 System Architecture', 26),
        ('   4.2 Technology Stack', 27),
        ('   4.3 Database Design', 28),
        ('5. Results and Analysis', 32),
        ('   5.1 Model Performance Comparison', 32),
        ('   5.2 Feature Importance', 33),
        ('   5.3 Discussion', 35),
        ('6. Web Application', 36),
        ('7. Conclusion and Future Work', 38),
        ('References', 42),
        ('Appendices', 43),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        tab_stops = p.paragraph_format.tab_stops
        p.add_run('\t' + str(page))
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 1: INTRODUCTION
    # ============================================
    
    doc.add_heading('Chapter 1: Introduction', level=1)
    
    doc.add_heading('1.1 Background and Motivation', level=2)
    doc.add_paragraph("""Agriculture is the backbone of the Indian economy, employing over 50% of the workforce and contributing significantly to the GDP. However, farmers in India face numerous challenges, with price volatility being one of the most critical issues affecting their economic stability. The inability to predict market prices leads to:

• Distress selling at low prices during harvest seasons
• Inadequate planning for crop selection
• Financial losses due to price fluctuations
• Reduced investment in agricultural inputs

West Bengal, one of India's leading agricultural states, produces significant quantities of rice, jute, and wheat. The state's diverse agro-climatic conditions and market dynamics create complex pricing patterns that are difficult to predict using traditional methods.""")
    
    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph("""The primary objective of this project is to develop an accurate and reliable commodity price prediction system that can:

1. Predict prices for major agricultural commodities (Rice, Jute, Wheat) in West Bengal
2. Provide 7-day ahead forecasts to aid decision-making
3. Incorporate multiple factors affecting prices including weather, economic indicators, and historical patterns
4. Offer an accessible web interface for farmers, traders, and policymakers""")
    
    doc.add_heading('1.3 Objectives', level=2)
    doc.add_paragraph("""The specific objectives of this BTP are:

1. Data Collection and Preprocessing: Compile comprehensive historical price data from 2014-2025, integrating economic and agricultural indicators

2. Feature Engineering: Design and implement relevant features capturing temporal, spatial, and economic patterns

3. Model Development: Implement and train two machine learning models:
   • XGBoost (Gradient Boosting)
   • Deep Neural Network with Backpropagation

4. Model Evaluation: Rigorously evaluate models using multiple metrics (MAPE, RMSE, R²)

5. Web Application Development: Create a user-friendly interface for price predictions

6. Deployment: Deploy the system using production-grade infrastructure""")
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 3: METHODOLOGY
    # ============================================
    
    doc.add_heading('Chapter 3: Methodology', level=1)
    
    doc.add_heading('3.1 Data Sources', level=2)
    
    # Data sources table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Source', 'Variables', 'Period']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    data = [
        ('Agmarknet', 'Daily modal prices', '2014-2025'),
        ('IMD', 'Temperature, Rainfall', '2014-2025'),
        ('RBI', 'CPI, Per Capita Income', '2014-2025'),
        ('Ministry of Agriculture', 'MSP, Area, Production, Yield', '2014-2025'),
        ('Ministry of Food', 'Food Subsidy', '2014-2025'),
        ('Fertilizer Association', 'Fertilizer Consumption', '2014-2025'),
        ('DGCIS', 'Export/Import data', '2014-2025'),
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Dataset Statistics', level=2)
    
    # Dataset table
    table2 = doc.add_table(rows=8, cols=2)
    table2.style = 'Table Grid'
    
    stats = [
        ('Parameter', 'Value'),
        ('Total Records', '177,320'),
        ('Time Period', '2014-01-01 to 2025-10-28'),
        ('Districts', '18'),
        ('Markets', '61'),
        ('Commodities', '3 (Rice, Jute, Wheat)'),
        ('Varieties', '14'),
        ('Database Size', '51.14 MB (SQLite)'),
    ]
    
    for i, (param, value) in enumerate(stats):
        table2.rows[i].cells[0].text = param
        table2.rows[i].cells[1].text = value
        if i == 0:
            table2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            table2.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 5: RESULTS
    # ============================================
    
    doc.add_heading('Chapter 5: Results and Analysis', level=1)
    
    doc.add_heading('5.1 Model Performance Comparison', level=2)
    
    # Performance table
    perf_table = doc.add_table(rows=5, cols=4)
    perf_table.style = 'Table Grid'
    
    perf_data = [
        ('Metric', 'XGBoost', 'Neural Network', 'Best Model'),
        ('MAE (Rs)', '167.42', '143.73', 'Neural Network'),
        ('RMSE (Rs)', '285.36', '253.66', 'Neural Network'),
        ('MAPE (%)', '5.43', '4.64', 'Neural Network'),
        ('R² Score', '0.9453', '0.9601', 'Neural Network'),
    ]
    
    for i, row_data in enumerate(perf_data):
        for j, cell_data in enumerate(row_data):
            perf_table.rows[i].cells[j].text = cell_data
            if i == 0:
                perf_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 Neural Network Training Details', level=2)
    doc.add_paragraph("""The Neural Network was trained with the following configuration and results:

• Training Time: ~15 minutes (CPU mode)
• Epochs Completed: 82 (early stopping triggered)
• Best Epoch: 67 with val_loss = 0.04832
• Architecture: 256→128→64→32→1 with BatchNormalization and Dropout

Test Performance:
• Test RMSE: 253.66 Rs
• Test MAE: 143.73 Rs
• Test R²: 0.9601
• Test MAPE: 4.64%

2025 Prediction Performance:
• 2025 MAPE: 4.44%
• 2025 R²: 0.9724

Sample 2025 Predictions:
• North 24 Parganas, Wheat: Actual ₹2060 → Predicted ₹2105 (Error: 2.2%)
• Birbhum, Wheat: Actual ₹2080 → Predicted ₹2072 (Error: 0.4%)
• Murshidabad, Jute: Actual ₹5880 → Predicted ₹5894 (Error: 0.2%)
• Medinipur(W), Rice: Actual ₹3800 → Predicted ₹3811 (Error: 0.3%)""")
    
    doc.add_heading('5.3 Commodity-wise Performance', level=2)
    
    # Commodity MAPE table
    comm_table = doc.add_table(rows=5, cols=3)
    comm_table.style = 'Table Grid'
    
    comm_data = [
        ('Commodity', 'XGBoost MAPE', 'Neural Network MAPE'),
        ('Rice', '5.21%', '4.42%'),
        ('Jute', '5.89%', '4.95%'),
        ('Wheat', '5.45%', '4.68%'),
        ('Overall', '5.43%', '4.64%'),
    ]
    
    for i, row_data in enumerate(comm_data):
        for j, cell_data in enumerate(row_data):
            comm_table.rows[i].cells[j].text = cell_data
            if i == 0 or i == 4:
                comm_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ============================================
    # CHAPTER 7: CONCLUSION
    # ============================================
    
    doc.add_heading('Chapter 7: Conclusion and Future Work', level=1)
    
    doc.add_heading('7.1 Summary', level=2)
    doc.add_paragraph("""This BTP project successfully developed a comprehensive commodity price prediction system for West Bengal, India. The key achievements are:

1. Data Integration: Created a comprehensive dataset of 177,320 records spanning 11 years (2014-2025), integrating price data with economic and agricultural indicators

2. Feature Engineering: Designed 30 meaningful features capturing temporal patterns, market dynamics, and economic factors

3. Model Development: Implemented two machine learning models:
   • XGBoost achieving 5.43% MAPE and 0.9453 R²
   • Neural Network achieving 4.64% MAPE and 0.9601 R²
   • Neural Network 2025 predictions: 4.44% MAPE and 0.9724 R²

4. Web Application: Developed a user-friendly React-based interface for accessing predictions

5. Production Deployment: Deployed using Waitress WSGI server for reliable operation""")
    
    doc.add_heading('7.2 Key Findings', level=2)
    doc.add_paragraph("""After retraining with enhanced 2024-2025 data, the Neural Network outperforms XGBoost:

1. Better Generalization: The Neural Network achieves lower MAPE (4.64% vs 5.43%) indicating better generalization

2. Excellent 2025 Predictions: For 2025 data specifically, the NN achieves remarkable R² = 0.9724 and MAPE = 4.44%

3. Temporal Pattern Capture: Deep learning effectively captures seasonal and temporal patterns in price data

4. Feature Engineering: 30 carefully engineered features including cyclical encodings for time contribute to accuracy""")
    
    doc.add_heading('7.3 Future Work', level=2)
    doc.add_paragraph("""Short-term Improvements:
• LSTM Integration: Implement Long Short-Term Memory networks for better temporal pattern capture
• Ensemble Methods: Combine XGBoost and Neural Network predictions
• Real-time Data: Integrate live weather and market data APIs

Long-term Goals:
• Geographic Expansion: Extend to other Indian states
• Commodity Expansion: Include vegetables, fruits, and other crops
• Mobile Application: Develop native mobile apps for easier access
• Multi-language Support: Add regional language interfaces""")
    
    # Save document
    output_path = 'BTP_Report.docx'
    doc.save(output_path)
    print(f"✅ Word document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_word_report()
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call(['pip', 'install', 'python-docx'])
        create_word_report()
